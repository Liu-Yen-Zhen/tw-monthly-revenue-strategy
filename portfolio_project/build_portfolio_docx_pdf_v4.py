#!/usr/bin/env python3
"""Build V4 industry-style Taiwan monthly-revenue research report Word/PDF.

V4 expands narrative into a professional research report: research question,
methodology, empirical interpretation, due-diligence risks, paper-trading plan,
and interview talking guide. Research-only; no trading or deployment.
"""
from __future__ import annotations

import csv
import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import Image, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle, HRFlowable

ROOT = Path('/Users/liuyenzhen/quant-research/tw_monthly_revenue')
PORT = ROOT / 'portfolio_project'
CHARTS = ROOT / 'reports' / 'charts'
V3CHARTS = PORT / 'charts_v3'
PROCESSED = ROOT / 'data' / 'processed'

RESUME_DOCX = PORT / 'Taiwan_Monthly_Revenue_Strategy_Resume_Version_ZH_PRO_V4.docx'
RESUME_PDF = PORT / 'Taiwan_Monthly_Revenue_Strategy_Resume_Version_ZH_PRO_V4.pdf'
GUIDE_DOCX = PORT / 'Taiwan_Monthly_Revenue_Strategy_Talking_Guide_ZH_PRO_V4.docx'
GUIDE_PDF = PORT / 'Taiwan_Monthly_Revenue_Strategy_Talking_Guide_ZH_PRO_V4.pdf'

pdfmetrics.registerFont(UnicodeCIDFont('STSong-Light'))
FONT = 'STSong-Light'
NAVY = colors.HexColor('#0F172A')
BLUE = colors.HexColor('#1D4ED8')
GREEN = colors.HexColor('#047857')
AMBER = colors.HexColor('#B45309')
RED = colors.HexColor('#B91C1C')
MUTED = colors.HexColor('#64748B')
LIGHT = colors.HexColor('#F8FAFC')
BORDER = colors.HexColor('#CBD5E1')
WHITE = colors.white


def read_csv(path: Path) -> list[dict[str, Any]]:
    if not path.exists():
        return []
    with path.open(encoding='utf-8') as f:
        return list(csv.DictReader(f))


def f(x: Any) -> float:
    try:
        return float(x)
    except Exception:
        return 0.0


def pct(x: Any) -> str:
    return f'{f(x):.1%}'


def num(x: Any) -> str:
    return f'{f(x):.2f}'


def charts() -> list[tuple[str, str, str]]:
    items = [
        (CHARTS/'s1_nav_drawdown_zh.png', 'Core performance', 'S1 NAV 與 Drawdown：觀察策略累積績效、回撤深度與修復能力；這是 proxy backtest，不代表實盤可複製績效。'),
        (CHARTS/'s1_monthly_returns_is_oos_zh.png', 'IS/OOS behavior', '月報酬與 IS/OOS 切分：檢查績效是否集中在少數月份，以及 2025 test 是否仍有正貢獻。'),
        (CHARTS/'signal_search_sharpe_mdd_scatter_zh.png', 'Search discipline', '策略搜尋 Sharpe vs MDD：業界評估不只看最高 Sharpe，也要同時看 drawdown、穩健性與可解釋性。'),
        (CHARTS/'top_variants_sharpe_vs_remove5_zh.png', 'Winner concentration', 'Top variants vs remove-top-5：檢查高績效版本是否由少數大贏家支撐。'),
        (CHARTS/'remove_winners_sharpe_decay_zh.png', 'Tail dependence', 'Remove-winners Sharpe decay：right-tail dependence 是目前最需要揭露與管理的風險。'),
        (CHARTS/'phase3_12_walkforward_oos_nav_zh.png', 'Walk-forward OOS', 'Walk-forward OOS NAV：train-selected rule 未穩定擊敗固定 S1，支持保留簡潔 incumbent 而非過度調參。'),
        (CHARTS/'phase3_12_sector_survival_zh.png', 'Sector survival', 'Sector survival：電子 / 半導體供應鏈解釋較強，no-semiconductor 明顯轉弱，不能宣稱 broad-market alpha。'),
        (CHARTS/'phase3_17_quiet_digestion_nav_zh.png', 'Price-volume diagnostic', 'Quiet digestion：作為市場消化基本面 surprise 的狀態變數有研究價值，但 standalone 交易數少。'),
        (CHARTS/'phase3_18_dynamic_sizing_nav_zh.png', 'Sizing overlay', 'Dynamic sizing：quiet boost 小幅改善，但未改善 drawdown 到足以 promotion。'),
        (V3CHARTS/'phase4_1_frr_top_sharpe_zh.png', 'New FRR strategy', 'FRR Top Sharpe：融資清洗反彈有訊號，但最佳 Sharpe 仍低於 S1，且交易數偏少。'),
        (V3CHARTS/'phase4_1_frr_remove_winner_zh.png', 'FRR robustness', 'FRR remove-winner stress：移除 top winners 後快速轉弱，不能作為 S1 替代策略。'),
        (V3CHARTS/'phase4_1_frr_sector_zh.png', 'FRR sector behavior', 'FRR sector survival：部分非半導體訊號存在，但樣本不足，適合作為 timing / sizing diagnostic。'),
    ]
    return [(str(p), title, cap) for p, title, cap in items if p.exists()]


def collect_tables() -> dict[str, list[list[str]]]:
    frr_variants = sorted(read_csv(PROCESSED / 'frr_margin_deleveraging_variants.csv'), key=lambda r: f(r.get('sharpe_cash_counted')), reverse=True)[:8]
    frr_top = [[r['variant'].replace('_', ' '), r['delay_trading_days'], r['holding_days'], pct(r['total_return']), num(r['sharpe_cash_counted']), pct(r['mdd']), r['trades']] for r in frr_variants]
    frr_rem = read_csv(PROCESSED / 'frr_margin_deleveraging_remove_winners.csv')
    frr_rem_rows = [[r['variant'].replace('_', ' '), r['remove_top_n'], pct(r['total_return']), num(r['sharpe_cash_counted']), pct(r['mdd'])] for r in frr_rem if r.get('variant') in {'frr2_no_catch_falling_knife','frr3_volume_absorption'}]
    return {
        'headline': [
            ['S1 incumbent', '+167.5%', '≈2.40', '-7.9%', '研究候選主版本；仍需 exact timing / paper trading'],
            ['S1 fixed-20 comparator', '+161.9%', '1.55', '-21.2%', '最簡潔 benchmark；用來隔離 exit-rule 貢獻'],
            ['Quiet boost sizing', '+174.4%', '1.62', '-21.2%', '小幅改善但 drawdown 未改善，未升級'],
            ['Conservative execution proxy', '+111.9%', '1.24', '-24.9%', 'next-open + 1.0% cost + non-fill stress 後 quality 下降'],
            ['FRR best first-pass', '+251.4%', '1.55', '-17.6%', '新融資策略；樣本少且右尾依賴，僅診斷候選'],
        ],
        'methodology': [
            ['Research question', '月營收 surprise 是否造成 10–20D post-disclosure repricing？', '從市場制度與資訊反應速度出發，而非指標 mining。'],
            ['Signal construction', '3M SUR persistence + not-overheated momentum', '捕捉持續性 surprise，同時避免已過度 pricing 的標的。'],
            ['Portfolio construction', 'Top 8, industry cap = 3, liquidity ≥ 50m TWD', '降低集中度與不可交易小型股假象。'],
            ['Execution assumption', 'next-open / delayed-entry / 1.0% cost stress', '避免用 same-day 或 close proxy 過度高估。'],
            ['Promotion gates', 'OOS, remove-winner, sector, cost, liquidity, timing, paper trading', '業界式 due diligence，不以單一 Sharpe 決策。'],
        ],
        'risk': [
            ['Data timing risk', '缺完整公司級公告 timestamp', '不能宣稱 same-day tradability；需補精確公告時間。'],
            ['Survivorship risk', '歷史 universe / 下市資料仍需加強', '目前結果可能受樣本可得性影響。'],
            ['Sector concentration', '電子 / 半導體供應鏈依賴明顯', '策略定位應限縮為 supply-chain repricing candidate。'],
            ['Right-tail dependence', 'remove-winner 後 Sharpe 衰退', '需要 sizing、capacity 與 risk budget 管控。'],
            ['Execution risk', '漲停不可成交、滑價、開盤跳空', '紙上績效可能高於實際可執行績效。'],
            ['Overfitting risk', '多輪參數搜尋容易追逐樣本內 winners', '保留簡潔 S1，不盲目升級複雜版本。'],
        ],
        'paper': [
            ['Daily process', '盤後更新營收 / 市場 / 融資資料，生成 next-day candidate list', '檢查資料真的能在預期時間取得。'],
            ['Timestamp log', 'data_observed_at, signal_generated_at, planned_entry_date', '驗證 anti-look-ahead 與 operational workflow。'],
            ['Execution log', 'planned open/close entry, actual tradability, non-fill reason, slippage', '量化 non-fill 與滑價偏差。'],
            ['Outcome review', '5D / 10D / 20D return, benchmark excess, assumption drift', '比較 paper trade 與 backtest 假設。'],
            ['Monthly review', 'winner concentration, sector exposure, liquidity, cost drift', '決定是否進入下一個 research gate。'],
        ],
        'frr_top': frr_top,
        'frr_remove': frr_rem_rows,
        'roadmap': [
            ['1', '補公司級 exact announcement timestamp', '判斷真正 earliest tradable date。'],
            ['2', '加強 historical universe / survivorship control', '提高研究可信度。'],
            ['3', '建立 auction / limit-up fill realism', '處理台股漲跌停與開盤跳空。'],
            ['4', '3–6 個月 paper trading', '驗證資料更新、可成交性、滑價與策略穩定性。'],
            ['5', '測 FRR as S1 sizing diagnostic', '把融資清洗用於加減碼，而非獨立 promotion。'],
        ],
        'qa': [
            ['這個策略的 alpha 來源是什麼？', '來自台灣月營收制度造成的高頻基本面資訊更新，以及市場對持續性 surprise 的延遲反應。'],
            ['為什麼不是技術指標？', '訊號從公告制度、營收 surprise、預期修正與資金逐步 repricing 建立；price-volume 只是後續消化狀態診斷。'],
            ['最大弱點是什麼？', 'exact timing、sector concentration、winner dependence、execution realism；文件中都明確揭露。'],
            ['為什麼 S1 而不是更複雜版本？', 'walk-forward 沒穩定擊敗 S1，複雜版本容易 overfit；業界會偏好簡潔且可解釋的 incumbent。'],
            ['FRR 新策略怎麼定位？', 'FRR 有因果邏輯但 remove-winner 脆弱，暫時作為 S1 timing / sizing diagnostic，不取代 S1。'],
        ],
    }

# ---------- DOCX helpers ----------
def set_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)


def style_doc(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Inches(.6); sec.bottom_margin = Inches(.55); sec.left_margin = Inches(.7); sec.right_margin = Inches(.7)
    styles = doc.styles
    for name in ['Normal', 'Heading 1', 'Heading 2', 'Heading 3']:
        styles[name].font.name = 'Microsoft JhengHei'
        styles[name]._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft JhengHei')
    styles['Normal'].font.size = Pt(10)
    styles['Heading 1'].font.size = Pt(17); styles['Heading 1'].font.color.rgb = RGBColor(15,23,42)
    styles['Heading 2'].font.size = Pt(12.5); styles['Heading 2'].font.color.rgb = RGBColor(29,78,216)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers)); table.style = 'Table Grid'
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]; c.text = h; set_shading(c, '0F172A')
        for p in c.paragraphs:
            for r in p.runs:
                r.font.bold = True; r.font.color.rgb = RGBColor(255,255,255)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row): cells[i].text = str(val)
    doc.add_paragraph('')


def add_chart(doc: Document, path: str, caption: str) -> None:
    p = Path(path)
    if p.exists(): doc.add_picture(str(p), width=Inches(6.6))
    cap = doc.add_paragraph(caption); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in cap.runs:
        r.font.size = Pt(8); r.font.color.rgb = RGBColor(100,116,139)


def add_doc_content(doc: Document, title: str, subtitle: str, is_guide: bool, ch: list[tuple[str,str,str]], tb: dict[str, list[list[str]]]) -> None:
    style_doc(doc)
    h = doc.add_heading(title, 0); h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(subtitle); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Research / Paper-Trading Only｜非投資建議｜非實盤交易系統').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    doc.add_heading('01｜Executive Summary：業界版研究摘要', 1)
    doc.add_paragraph('本研究檢驗台灣月營收公告制度是否能產生短週期、可解釋且可被風險管理的 post-disclosure repricing 機會。研究流程不是以技術指標或單一最佳化結果為起點，而是先定義市場結構、資料可得時間、訊號形成邏輯與可交易限制，再用 portfolio-level backtest 與多層 robustness gates 評估。')
    doc.add_paragraph('結論上，S1 目前可被定位為 portfolio-grade v0.1 research candidate：它有清楚的基本面 surprise 邏輯與吸引人的 proxy 結果，但仍未通過 production-ready 所需的 exact timestamp、survivorship、execution fill 與 paper-trading gates。')
    add_table(doc, ['版本', 'Return', 'Sharpe', 'MDD', '專業定位'], tb['headline'])

    doc.add_heading('02｜Research Question and Economic Rationale', 1)
    doc.add_paragraph('核心研究問題：若公司月營收出現連續性正向 surprise，市場是否會因資訊處理、產業鏈確認、法人調整持倉或散戶注意力延遲，而在公告後 10–20 個交易日內逐步重新定價？')
    doc.add_paragraph('因果鏈：月營收公告制度 → 高頻基本面資訊更新 → 3M SUR persistence → 投資人預期修正 → 短週期 repricing。這個 framing 的重點是「誰因為什麼制度或行為限制而反應不足」，而不是事後挑選 K 線形狀。')
    add_table(doc, ['研究模組', '設計', '業界解讀'], tb['methodology'])

    doc.add_heading('03｜Data, Timing Controls, and Backtest Design', 1)
    doc.add_paragraph('資料使用官方月營收、每日價格 / 成交金額、官方 daily raw JSON、OHLC / limit parser，以及後續新增的官方融資餘額。所有資料都必須區分 event month、data-available date、signal date、trade date。由於目前歷史月營收仍缺完整公司級公告 timestamp，本報告不宣稱 same-day tradability。')
    doc.add_paragraph('Portfolio construction 採用 Top 8、industry cap、liquidity threshold 與成本假設，避免績效只來自不可交易小型股或單一產業集中。策略 promotion 需要同時通過年分切分、walk-forward、remove-winner、sector survival、流動性與 execution realism，而不是只看 full-sample Sharpe。')

    doc.add_heading('04｜Empirical Findings and Interpretation', 1)
    for i, (path, _title, cap) in enumerate(ch[:9], 1):
        add_chart(doc, path, cap)
        if i in {2,5,7}: doc.add_page_break()
    doc.add_paragraph('整體來看，S1 的吸引力來自「結果、解釋性與簡潔度」的平衡；更複雜的 quiet digestion / dynamic sizing 版本雖然局部改善報酬，但未能穩定改善 drawdown 或 OOS robustness，因此不應取代 incumbent。')

    doc.add_heading('05｜FRR 融資清洗反彈：新增策略研究結論', 1)
    doc.add_paragraph('FRR 測試另一條市場結構邏輯：強基本面股票若經歷融資去槓桿，短期槓桿賣壓釋放後可能反彈。第一輪結果顯示「融資下降 + 放量換手」比單純不破低更有訊號，但樣本少、remove-winner 脆弱、高流動性版本弱化，因此目前只能保留為 S1 的 timing / sizing diagnostic。')
    add_table(doc, ['Variant','Delay','Hold','Return','Sharpe','MDD','Trades'], tb['frr_top'])
    for path, _title, cap in ch[9:]: add_chart(doc, path, cap)
    add_table(doc, ['Variant','Remove top N','Return','Sharpe','MDD'], tb['frr_remove'][:8])

    doc.add_heading('06｜Due Diligence: Risks and Non-promotion Rationale', 1)
    doc.add_paragraph('業界報告最重要的是把限制講清楚。這個專案的價值不在於宣稱策略已可實盤，而在於展示完整 research-to-validation workflow，並且能誠實拒絕沒有通過 gate 的版本。')
    add_table(doc, ['風險', '目前狀態', '對研究結論的影響'], tb['risk'])

    doc.add_heading('07｜Paper Trading and Production Gate Plan', 1)
    doc.add_paragraph('下一階段不應繼續盲目擴大參數搜尋，而應轉向 paper-trading validation。Paper trading 的目的是驗證資料更新時間、訊號生成時間、實際可成交性、滑價、non-fill 與回測假設是否一致。')
    add_table(doc, ['流程', '記錄欄位 / 動作', '驗證目的'], tb['paper'])
    add_table(doc, ['Gate', '下一步', '目的'], tb['roadmap'])

    doc.add_heading('08｜Final Positioning', 1)
    doc.add_paragraph('建議對外定位：這是一個台股月營收 surprise 的 portfolio-grade quant research project，展示從市場制度假說、官方資料 pipeline、SUR-style factor design、portfolio backtest、robustness gates、execution realism 到 paper-trading schema 的完整研究流程。')
    doc.add_paragraph('避免說法：production-ready alpha、live trading system、保證 Sharpe 2+、全市場通用台股策略。')

    if is_guide:
        doc.add_page_break()
        doc.add_heading('09｜Interview Talking Guide', 1)
        doc.add_heading('30 秒 pitch', 2)
        doc.add_paragraph('我做的是台股月營收 surprise 的短週期量化研究。核心從台灣每月公布營收的制度出發，測試持續性營收驚喜是否會在 10–20 個交易日內被市場逐步重估。研究中我不只看 Sharpe，也做 OOS、sector survival、remove-winner、execution timing、融資策略延伸與 paper-trading schema，所以最後我把 S1 定位為 portfolio-grade research candidate，而不是 production-ready alpha。')
        doc.add_heading('3 分鐘 pitch', 2)
        doc.add_paragraph('第一，市場結構：台灣每月營收提供高頻基本面資訊。第二，訊號設計：用 3M SUR persistence 捕捉持續 surprise，並避開已過熱 momentum。第三，portfolio 與風控：Top 8、industry cap、liquidity gate、成本與延遲進場。第四，robustness：walk-forward、sector survival、remove winners、execution realism。第五，結論：S1 有研究價值但仍需 exact timestamp 與 paper trading；FRR 融資策略有診斷價值但不升級。')
        add_table(doc, ['面試問題', '專業回答'], tb['qa'])

# ---------- PDF helpers ----------
def styles():
    s = getSampleStyleSheet()
    s.add(ParagraphStyle('TitleC', fontName=FONT, fontSize=23, leading=30, textColor=NAVY, alignment=TA_CENTER, spaceAfter=9))
    s.add(ParagraphStyle('SubC', fontName=FONT, fontSize=10.5, leading=15, textColor=MUTED, alignment=TA_CENTER, spaceAfter=7))
    s.add(ParagraphStyle('H1C', fontName=FONT, fontSize=14.2, leading=19, textColor=NAVY, spaceBefore=8, spaceAfter=5))
    s.add(ParagraphStyle('H2C', fontName=FONT, fontSize=10.5, leading=14, textColor=BLUE, spaceBefore=5, spaceAfter=3))
    s.add(ParagraphStyle('BodyC', fontName=FONT, fontSize=8.7, leading=12.8, textColor=colors.HexColor('#111827'), wordWrap='CJK', spaceAfter=4))
    s.add(ParagraphStyle('SmallC', fontName=FONT, fontSize=7.2, leading=9.5, textColor=MUTED, wordWrap='CJK'))
    s.add(ParagraphStyle('CapC', fontName=FONT, fontSize=7.2, leading=9.5, textColor=MUTED, alignment=TA_CENTER, spaceAfter=6))
    return s
S = styles()

def esc(x: str) -> str: return str(x).replace('&','&amp;').replace('<','&lt;').replace('>','&gt;')
def P(text: str, style='BodyC') -> Paragraph: return Paragraph(esc(text), S[style])
def section(title: str) -> list[Any]: return [Paragraph(esc(title), S['H1C']), HRFlowable(width='100%', thickness=.55, color=BORDER, spaceAfter=5)]

def pdf_table(headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> Table:
    data = [[Paragraph(f'<b>{esc(h)}</b>', S['SmallC']) for h in headers]] + [[Paragraph(esc(c), S['SmallC']) for c in r] for r in rows]
    if widths is None: widths = [16.9/len(headers)*cm]*len(headers)
    t = Table(data, colWidths=widths, repeatRows=1)
    t.setStyle(TableStyle([('BACKGROUND',(0,0),(-1,0),NAVY),('TEXTCOLOR',(0,0),(-1,0),WHITE),('ROWBACKGROUNDS',(0,1),(-1,-1),[WHITE,LIGHT]),('BOX',(0,0),(-1,-1),.5,BORDER),('INNERGRID',(0,0),(-1,-1),.3,BORDER),('VALIGN',(0,0),(-1,-1),'TOP'),('LEFTPADDING',(0,0),(-1,-1),4),('RIGHTPADDING',(0,0),(-1,-1),4),('TOPPADDING',(0,0),(-1,-1),4),('BOTTOMPADDING',(0,0),(-1,-1),4)]))
    return t

def pdf_chart(path: str, caption: str, max_h=7.6*cm) -> list[Any]:
    p=Path(path)
    if not p.exists(): return [P(f'缺圖：{path}', 'SmallC')]
    img=Image(str(p)); scale=min(16.8*cm/img.imageWidth, max_h/img.imageHeight, 1)
    img.drawWidth=img.imageWidth*scale; img.drawHeight=img.imageHeight*scale
    return [img, Paragraph(esc(caption), S['CapC'])]

def footer(canvas, doc):
    canvas.saveState(); canvas.setFont(FONT, 7); canvas.setFillColor(MUTED)
    canvas.drawString(1.6*cm, .72*cm, 'Taiwan Monthly Revenue Strategy｜Professional Research Report｜Research Only')
    canvas.drawRightString(A4[0]-1.6*cm, .72*cm, f'Page {doc.page}')
    canvas.restoreState()

def build_pdf(path: Path, title: str, subtitle: str, is_guide: bool, ch: list[tuple[str,str,str]], tb: dict[str, list[list[str]]]) -> None:
    story: list[Any] = [Spacer(1,1.8*cm), Paragraph(esc(title), S['TitleC']), Paragraph(esc(subtitle), S['SubC']), pdf_table(['文件定位','狀態'], [[('履歷附件 / 業界研究報告版' if not is_guide else '面試講稿 / 業界答辯版'), 'Research / Paper-Trading Only；非投資建議；非實盤交易系統']], [7*cm,9.9*cm]), PageBreak()]
    story += section('01｜Executive Summary：業界版研究摘要')
    story += [P('本研究檢驗台灣月營收公告制度是否能產生短週期、可解釋且可被風險管理的 post-disclosure repricing 機會。研究流程不是以技術指標或單一最佳化結果為起點，而是先定義市場結構、資料可得時間、訊號形成邏輯與可交易限制，再用 portfolio-level backtest 與多層 robustness gates 評估。'), P('結論上，S1 目前可被定位為 portfolio-grade v0.1 research candidate：它有清楚的基本面 surprise 邏輯與吸引人的 proxy 結果，但仍未通過 production-ready 所需的 exact timestamp、survivorship、execution fill 與 paper-trading gates。'), pdf_table(['版本','Return','Sharpe','MDD','專業定位'], tb['headline'], [3.4*cm,2*cm,1.8*cm,1.8*cm,7.9*cm]), PageBreak()]
    story += section('02｜Research Question and Methodology')
    story += [P('核心研究問題：若公司月營收出現連續性正向 surprise，市場是否會因資訊處理、產業鏈確認、法人調整持倉或散戶注意力延遲，而在公告後 10–20 個交易日內逐步重新定價？'), P('因果鏈：月營收公告制度 → 高頻基本面資訊更新 → 3M SUR persistence → 投資人預期修正 → 短週期 repricing。'), pdf_table(['研究模組','設計','業界解讀'], tb['methodology'], [3.2*cm,5.8*cm,7.9*cm]), PageBreak()]
    story += section('03｜Data, Timing Controls, and Promotion Gates')
    story += [P('資料使用官方月營收、每日價格 / 成交金額、官方 daily raw JSON、OHLC / limit parser，以及官方融資餘額。所有資料都必須區分 event month、data-available date、signal date、trade date。由於目前歷史月營收仍缺完整公司級公告 timestamp，本報告不宣稱 same-day tradability。'), pdf_table(['風險','目前狀態','對研究結論的影響'], tb['risk'], [3.4*cm,6.1*cm,7.4*cm]), PageBreak()]
    story += section('04｜Empirical Findings: Core Charts')
    for i,(p,_t,cap) in enumerate(ch[:2],1): story += pdf_chart(p, cap, 7.8*cm)
    story += [PageBreak()] + section('05｜Robustness and Search Discipline')
    for p,_t,cap in ch[2:5]: story += pdf_chart(p, cap, 5.9*cm)
    story += [PageBreak()] + section('06｜Walk-forward, Sector Survival, and Extensions')
    for p,_t,cap in ch[5:9]: story += pdf_chart(p, cap, 5.25*cm)
    story += [PageBreak()] + section('07｜FRR Margin-deleveraging Strategy')
    story += [P('FRR 測試另一條市場結構邏輯：強基本面股票若經歷融資去槓桿，短期槓桿賣壓釋放後可能反彈。第一輪結果顯示「融資下降 + 放量換手」比單純不破低更有訊號，但樣本少、remove-winner 脆弱、高流動性版本弱化，因此目前只能保留為 S1 的 timing / sizing diagnostic。'), pdf_table(['Variant','Delay','Hold','Return','Sharpe','MDD','Trades'], tb['frr_top'], [4.6*cm,1.3*cm,1.3*cm,2*cm,1.8*cm,2*cm,1.5*cm])]
    for p,_t,cap in ch[9:]: story += pdf_chart(p, cap, 5.4*cm)
    story += [PageBreak()] + section('08｜FRR Stress Test and Paper-trading Plan')
    story += [pdf_table(['Variant','Remove top N','Return','Sharpe','MDD'], tb['frr_remove'][:8], [5.6*cm,2.2*cm,2.3*cm,2.2*cm,2.3*cm]), Spacer(1,6), pdf_table(['流程','記錄欄位 / 動作','驗證目的'], tb['paper'], [3.4*cm,7.1*cm,6.4*cm]), PageBreak()]
    story += section('09｜Roadmap and Final Positioning')
    story += [pdf_table(['Gate','下一步','目的'], tb['roadmap'], [2.1*cm,6.7*cm,8.1*cm]), P('建議對外定位：這是一個台股月營收 surprise 的 portfolio-grade quant research project，展示從市場制度假說、官方資料 pipeline、SUR-style factor design、portfolio backtest、robustness gates、execution realism 到 paper-trading schema 的完整研究流程。避免宣稱 production-ready alpha、live trading system、保證 Sharpe 2+ 或全市場通用台股策略。')]
    if is_guide:
        story += [PageBreak()] + section('10｜Interview Talking Guide')
        story += [Paragraph('30 秒 pitch', S['H2C']), P('我做的是台股月營收 surprise 的短週期量化研究。核心從台灣每月公布營收的制度出發，測試持續性營收驚喜是否會在 10–20 個交易日內被市場逐步重估。研究中我不只看 Sharpe，也做 OOS、sector survival、remove-winner、execution timing、融資策略延伸與 paper-trading schema，所以最後我把 S1 定位為 portfolio-grade research candidate，而不是 production-ready alpha。'), Paragraph('3 分鐘 pitch', S['H2C']), P('第一，市場結構：台灣每月營收提供高頻基本面資訊。第二，訊號設計：用 3M SUR persistence 捕捉持續 surprise，並避開已過熱 momentum。第三，portfolio 與風控：Top 8、industry cap、liquidity gate、成本與延遲進場。第四，robustness：walk-forward、sector survival、remove winners、execution realism。第五，結論：S1 有研究價值但仍需 exact timestamp 與 paper trading；FRR 融資策略有診斷價值但不升級。'), pdf_table(['面試問題','專業回答'], tb['qa'], [5.2*cm,11.7*cm])]
    doc = SimpleDocTemplate(str(path), pagesize=A4, leftMargin=1.6*cm, rightMargin=1.6*cm, topMargin=1.25*cm, bottomMargin=1.2*cm)
    doc.build(story, onFirstPage=footer, onLaterPages=footer)

def main() -> int:
    ch = charts(); tb = collect_tables()
    d = Document(); add_doc_content(d, '台股月營收驚喜策略研究報告', '履歷附件版 V4｜業界專業敘述＋圖表強化版', False, ch, tb); d.save(RESUME_DOCX)
    g = Document(); add_doc_content(g, '台股月營收驚喜策略研究報告', '面試講稿版 V4｜業界答辯與圖表講解版', True, ch, tb); g.save(GUIDE_DOCX)
    build_pdf(RESUME_PDF, '台股月營收驚喜策略研究報告', '履歷附件版 V4｜業界專業敘述＋圖表強化版', False, ch, tb)
    build_pdf(GUIDE_PDF, '台股月營收驚喜策略研究報告', '面試講稿版 V4｜業界答辯與圖表講解版', True, ch, tb)
    for p in [RESUME_DOCX, RESUME_PDF, GUIDE_DOCX, GUIDE_PDF]: print(p, p.stat().st_size)
    return 0

if __name__ == '__main__':
    raise SystemExit(main())
