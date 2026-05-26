#!/usr/bin/env python3
"""Build enhanced Taiwan monthly-revenue portfolio Word/PDF deliverables.

Creates chart-dense Traditional Chinese portfolio and talking-guide files.
Research-only; no trading, deployment, or package installation.
"""
from __future__ import annotations

import csv
import re
from pathlib import Path
from typing import Any

import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import Image, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle, HRFlowable

ROOT = Path('/Users/liuyenzhen/quant-research/tw_monthly_revenue')
PORT = ROOT / 'portfolio_project'
CHARTS = ROOT / 'reports' / 'charts'
OUTCHARTS = PORT / 'charts_v3'
PROCESSED = ROOT / 'data' / 'processed'

RESUME_DOCX = PORT / 'Taiwan_Monthly_Revenue_Strategy_Resume_Version_ZH_PRO_V3.docx'
RESUME_PDF = PORT / 'Taiwan_Monthly_Revenue_Strategy_Resume_Version_ZH_PRO_V3.pdf'
GUIDE_DOCX = PORT / 'Taiwan_Monthly_Revenue_Strategy_Talking_Guide_ZH_PRO_V3.docx'
GUIDE_PDF = PORT / 'Taiwan_Monthly_Revenue_Strategy_Talking_Guide_ZH_PRO_V3.pdf'

pdfmetrics.registerFont(UnicodeCIDFont('STSong-Light'))
FONT = 'STSong-Light'
NAVY = colors.HexColor('#0F172A')
BLUE = colors.HexColor('#1D4ED8')
GREEN = colors.HexColor('#047857')
AMBER = colors.HexColor('#B45309')
RED = colors.HexColor('#B91C1C')
MUTED = colors.HexColor('#64748B')
LIGHT = colors.HexColor('#F8FAFC')
BORDER = colors.HexColor('#CBD5E1')
WHITE = colors.white


def read_csv(path: Path) -> list[dict[str, Any]]:
    if not path.exists():
        return []
    with path.open(encoding='utf-8') as f:
        return list(csv.DictReader(f))


def f(x: Any) -> float:
    try:
        return float(x)
    except Exception:
        return 0.0


def pct(x: Any) -> str:
    return f'{f(x):.1%}'


def num(x: Any) -> str:
    return f'{f(x):.2f}'


def setup_matplotlib() -> None:
    candidates = ['PingFang TC', 'Heiti TC', 'Arial Unicode MS', 'Microsoft JhengHei', 'Noto Sans CJK TC', 'DejaVu Sans']
    plt.rcParams['font.sans-serif'] = candidates
    plt.rcParams['axes.unicode_minus'] = False


def make_extra_charts() -> list[tuple[str, str]]:
    OUTCHARTS.mkdir(parents=True, exist_ok=True)
    setup_matplotlib()
    made: list[tuple[str, str]] = []

    variants = read_csv(PROCESSED / 'frr_margin_deleveraging_variants.csv')
    top = sorted(variants, key=lambda r: f(r.get('sharpe_cash_counted')), reverse=True)[:10]
    if top:
        labels = [f"{r['variant'].replace('frr','FRR-')[:18]}\nD{r['delay_trading_days']} H{r['holding_days']}" for r in top]
        vals = [f(r['sharpe_cash_counted']) for r in top]
        fig, ax = plt.subplots(figsize=(12, 6.2))
        bars = ax.bar(range(len(vals)), vals, color='#2563EB')
        ax.axhline(2.40, color='#B91C1C', linestyle='--', linewidth=1.3, label='S1 incumbent Sharpe ≈ 2.40')
        ax.set_title('FRR 融資清洗策略搜尋：Top 10 Sharpe（仍低於 S1）')
        ax.set_ylabel('Cash-counted Sharpe')
        ax.set_xticks(range(len(labels)), labels, rotation=35, ha='right')
        ax.legend()
        ax.grid(axis='y', alpha=.25)
        for b, v in zip(bars, vals):
            ax.text(b.get_x()+b.get_width()/2, v+0.03, f'{v:.2f}', ha='center', fontsize=9)
        fig.tight_layout()
        p = OUTCHARTS / 'phase4_1_frr_top_sharpe_zh.png'
        fig.savefig(p, dpi=180); plt.close(fig)
        made.append((str(p), '圖 10｜FRR 策略搜尋 Top 10：最佳 Sharpe 約 1.55，仍低於 S1 incumbent，適合作為診斷層而非替代策略。'))

    rem = read_csv(PROCESSED / 'frr_margin_deleveraging_remove_winners.csv')
    if rem:
        fig, ax = plt.subplots(figsize=(10.5, 5.8))
        for variant, color in [('frr1_basic_deleveraging', '#64748B'), ('frr2_no_catch_falling_knife', '#2563EB'), ('frr3_volume_absorption', '#047857')]:
            rows = sorted([r for r in rem if r['variant'] == variant], key=lambda r: int(r['remove_top_n']))
            if rows:
                ax.plot([int(r['remove_top_n']) for r in rows], [f(r['sharpe_cash_counted']) for r in rows], marker='o', label=variant.replace('_', ' '), color=color)
        ax.axhline(0, color='#111827', linewidth=.8)
        ax.set_title('FRR remove-winner stress：移除大贏家後 Sharpe 快速衰退')
        ax.set_xlabel('Remove top N winning trades')
        ax.set_ylabel('Sharpe')
        ax.legend(fontsize=8)
        ax.grid(alpha=.25)
        fig.tight_layout()
        p = OUTCHARTS / 'phase4_1_frr_remove_winner_zh.png'
        fig.savefig(p, dpi=180); plt.close(fig)
        made.append((str(p), '圖 11｜FRR remove-winner stress：FRR-2 / FRR-3 移除 top 10 後轉負，顯示右尾依賴仍是主要限制。'))

    sector = read_csv(PROCESSED / 'frr_margin_deleveraging_sector.csv')
    rows = [r for r in sector if r.get('variant') == 'frr3_volume_absorption']
    if rows:
        order = ['all', 'electronics', 'non_electronics', 'semiconductor', 'no_semiconductor']
        rows = sorted(rows, key=lambda r: order.index(r['subset']) if r['subset'] in order else 99)
        fig, ax = plt.subplots(figsize=(10.5, 5.8))
        vals = [f(r['sharpe_cash_counted']) for r in rows]
        labels = [r['subset'] for r in rows]
        colors_ = ['#2563EB', '#0284C7', '#047857', '#B45309', '#7C3AED']
        bars = ax.bar(labels, vals, color=colors_[:len(vals)])
        ax.set_title('FRR-3 sector survival：非半導體仍有訊號，但樣本數偏少')
        ax.set_ylabel('Sharpe')
        ax.grid(axis='y', alpha=.25)
        for b, v in zip(bars, vals):
            ax.text(b.get_x()+b.get_width()/2, v+0.03, f'{v:.2f}', ha='center', fontsize=10)
        fig.tight_layout()
        p = OUTCHARTS / 'phase4_1_frr_sector_zh.png'
        fig.savefig(p, dpi=180); plt.close(fig)
        made.append((str(p), '圖 12｜FRR-3 sector survival：非半導體 Sharpe 尚可，但交易數少，不能宣稱 broad alpha。'))
    return made


def existing_charts() -> list[tuple[str, str]]:
    items = [
        ('s1_nav_drawdown_zh.png', '圖 1｜S1 NAV 與 drawdown：展示績效路徑與回撤控制，但仍是 proxy backtest。'),
        ('s1_monthly_returns_is_oos_zh.png', '圖 2｜月報酬與 IS/OOS 切分：檢查樣本外是否仍有正貢獻。'),
        ('signal_search_sharpe_mdd_scatter_zh.png', '圖 3｜策略搜尋 Sharpe vs MDD：避免只追最高 Sharpe，需同時看回撤。'),
        ('top_variants_sharpe_vs_remove5_zh.png', '圖 4｜Top variants vs remove-top-5：檢查高績效是否依賴少數大贏家。'),
        ('remove_winners_sharpe_decay_zh.png', '圖 5｜Remove-winners Sharpe decay：right-tail dependence 是核心風險。'),
        ('phase3_12_walkforward_oos_nav_zh.png', '圖 6｜Walk-forward OOS NAV：固定 S1 未被 train-selected rule 穩定擊敗。'),
        ('phase3_12_sector_survival_zh.png', '圖 7｜Sector survival：電子 / 半導體依賴明顯，no-semiconductor 轉弱。'),
        ('phase3_17_quiet_digestion_nav_zh.png', '圖 8｜Quiet digestion NAV：具解釋性，但 standalone 樣本偏少。'),
        ('phase3_18_dynamic_sizing_nav_zh.png', '圖 9｜Dynamic sizing NAV：作為 sizing overlay 小幅改善，但不足以 promotion。'),
    ]
    return [(str(CHARTS / fn), cap) for fn, cap in items if (CHARTS / fn).exists()]


def collect_tables() -> dict[str, list[list[str]]]:
    frr_variants = sorted(read_csv(PROCESSED / 'frr_margin_deleveraging_variants.csv'), key=lambda r: f(r.get('sharpe_cash_counted')), reverse=True)[:8]
    frr_top = [[r['variant'].replace('_', ' '), r['delay_trading_days'], r['holding_days'], pct(r['total_return']), num(r['sharpe_cash_counted']), pct(r['mdd']), r['trades']] for r in frr_variants]
    frr_rem = read_csv(PROCESSED / 'frr_margin_deleveraging_remove_winners.csv')
    frr_rem_rows = [[r['variant'].replace('_', ' '), r['remove_top_n'], pct(r['total_return']), num(r['sharpe_cash_counted']), pct(r['mdd'])] for r in frr_rem if r.get('variant') in {'frr2_no_catch_falling_knife','frr3_volume_absorption'}]
    frr_sector = read_csv(PROCESSED / 'frr_margin_deleveraging_sector.csv')
    frr_sector_rows = [[r['variant'].replace('_', ' '), r['subset'], pct(r['total_return']), num(r['sharpe_cash_counted']), pct(r['mdd'])] for r in frr_sector if r.get('variant') in {'frr2_no_catch_falling_knife','frr3_volume_absorption'}]
    return {
        'headline': [
            ['S1 incumbent', '+167.5%', '≈2.40', '-7.9%', 'Portfolio-grade v0.1；仍需 exact timing / paper trading'],
            ['S1 fixed-20 comparator', '+161.9%', '1.55', '-21.2%', '簡潔 benchmark'],
            ['Quiet boost sizing', '+174.4%', '1.62', '-21.2%', '小幅改善，不 promotion'],
            ['Conservative execution proxy', '+111.9%', '1.24', '-24.9%', '保守假設後 quality 下降'],
            ['FRR best first-pass', '+251.4%', '1.55', '-17.6%', '新融資策略，樣本少、僅診斷候選'],
        ],
        'spec': [
            ['Signal', '3M SUR persistence', '降低單月營收雜訊，捕捉連續 surprise'],
            ['Filter', 'No overheated momentum', '避免股價已提前 pricing'],
            ['Liquidity', '20D avg turnover ≥ 50m TWD', '降低不可交易小型股假象'],
            ['Portfolio', 'Top 8, industry cap = 3', '控制單月與產業集中度'],
            ['Exit', '20D + sl8_trail12 proxy', '短週期 repricing + 風控 proxy'],
        ],
        'robustness': [
            ['Remove winners', 'Top 5/10/20 stress', '檢查右尾依賴'],
            ['Walk-forward', 'Train 2023→2024; 2023–24→2025', '檢查樣本外選參數是否有效'],
            ['Sector survival', 'electronics / semiconductor / no-semiconductor', '避免誤稱 broad alpha'],
            ['Execution realism', 'next-open, 1.0% cost, limit-up non-fill', '避免 close-price proxy 高估'],
            ['Paper trading', '資料觀測時間、non-fill、滑價、20D outcome', '驗證 operational feasibility'],
        ],
        'frr_top': frr_top,
        'frr_remove': frr_rem_rows,
        'frr_sector': frr_sector_rows,
        'roadmap': [
            ['Gate 1', '公司級 exact announcement timestamp', '釐清最早可交易時間'],
            ['Gate 2', 'survivorship / historical universe 補強', '降低樣本偏誤'],
            ['Gate 3', 'auction / limit-up fill realism', '驗證漲停不可成交與滑價'],
            ['Gate 4', 'paper trading 3–6 個月', '驗證流程與假設漂移'],
            ['Gate 5', 'FRR as S1 sizing diagnostic', '不獨立 promotion，先測是否改善 S1 sizing'],
        ]
    }


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)


def style_doc(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Inches(.55); sec.bottom_margin = Inches(.55); sec.left_margin = Inches(.65); sec.right_margin = Inches(.65)
    styles = doc.styles
    styles['Normal'].font.name = 'Microsoft JhengHei'; styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft JhengHei'); styles['Normal'].font.size = Pt(10)
    for s in ['Heading 1', 'Heading 2', 'Heading 3']:
        styles[s].font.name = 'Microsoft JhengHei'; styles[s]._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft JhengHei')
    styles['Heading 1'].font.size = Pt(18); styles['Heading 1'].font.color.rgb = RGBColor(15,23,42)
    styles['Heading 2'].font.size = Pt(13); styles['Heading 2'].font.color.rgb = RGBColor(29,78,216)


def docx_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], '0F172A')
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255,255,255); run.font.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    doc.add_paragraph('')


def docx_chart(doc: Document, path: str, caption: str) -> None:
    p = Path(path)
    if p.exists():
        doc.add_picture(str(p), width=Inches(6.7))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in cap.runs:
        r.font.size = Pt(8); r.font.color.rgb = RGBColor(100,116,139)


def add_docx_common(doc: Document, title: str, subtitle: str, is_guide: bool, charts: list[tuple[str, str]], tables: dict[str, list[list[str]]]) -> None:
    style_doc(doc)
    h = doc.add_heading(title, 0); h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(subtitle); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Research / Paper-Trading Only｜非投資建議｜非實盤交易系統').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    doc.add_heading('01｜Executive Snapshot', level=1)
    doc.add_paragraph('本作品集以台灣月營收公告制度為起點，建立 SUR-style 基本面 surprise 因子、短週期 portfolio backtest、robustness gates、execution realism 與 paper-trading schema。最終保留 S1 作為 portfolio-grade v0.1 研究候選；它不是 production-ready alpha。')
    docx_table(doc, ['版本', 'Return', 'Sharpe', 'MDD', '定位'], tables['headline'])

    doc.add_heading('02｜策略規格與因果鏈', level=1)
    doc.add_paragraph('因果鏈：月營收公告制度 → 持續性營收 surprise → 投資人預期逐步調整 → 公布後 10–20D repricing。')
    docx_table(doc, ['模組', '設定', '目的'], tables['spec'])

    doc.add_heading('03｜Data Pipeline & Anti-look-ahead', level=1)
    docx_table(doc, ['資料 / Gate', '做法', '目的'], tables['robustness'])
    doc.add_paragraph('重點限制：歷史資料尚未完整包含公司級 exact announcement timestamp，因此文件明確標示 next-open / delayed-entry / close-price proxy 的限制。')

    doc.add_heading('04｜核心圖表與結果解讀', level=1)
    for i, (path, cap) in enumerate(charts[:9], 1):
        docx_chart(doc, path, cap)
        if i in {2,5,7}:
            doc.add_page_break()

    doc.add_heading('05｜新策略 FRR：融資清洗反彈研究', level=1)
    doc.add_paragraph('FRR 測試「強營收 surprise + 融資去槓桿 + 放量換手 / 不接刀」是否形成短線反彈。第一輪結果顯示 headline 有訊號，但 Sharpe、交易數、remove-winner 與高流動性檢查不足以取代 S1。')
    docx_table(doc, ['Variant', 'Delay', 'Hold', 'Return', 'Sharpe', 'MDD', 'Trades'], tables['frr_top'])
    for path, cap in charts[9:]:
        docx_chart(doc, path, cap)
    docx_table(doc, ['Variant', 'Remove top N', 'Return', 'Sharpe', 'MDD'], tables['frr_remove'][:8])
    docx_table(doc, ['Variant', 'Subset', 'Return', 'Sharpe', 'MDD'], tables['frr_sector'])

    doc.add_heading('06｜限制、Roadmap 與結論定位', level=1)
    docx_table(doc, ['Gate', '下一步', '目的'], tables['roadmap'])
    doc.add_paragraph('最終定位：這個作品集展示完整量化研究流程，而不是宣稱已可實盤。最強敘述是：我能從市場制度提出假說、建立官方資料 pipeline、設計因子、做 portfolio-level robustness，並誠實拒絕未通過 production gates 的策略。')

    if is_guide:
        doc.add_page_break()
        doc.add_heading('Interview Talking Guide｜面試講稿', level=1)
        doc.add_heading('30 秒版本', level=2)
        doc.add_paragraph('我做的是台股月營收 surprise 的短週期量化研究。核心不是技術指標，而是利用台灣每月公布營收的制度，測試持續性營收驚喜是否會在 10–20 個交易日內被市場逐步重估。研究中我不只看 Sharpe，也做 OOS、sector survival、remove-winner、execution timing 與 paper-trading schema，所以最後我把 S1 定位為 portfolio-grade research candidate，而不是 production-ready alpha。')
        doc.add_heading('常見追問與回答', level=2)
        qas = [
            ['為什麼不用單月 YoY？', '單月 YoY 容易包含季節與一次性因素，所以我使用 3M SUR persistence 來捕捉較穩定的 surprise。'],
            ['最大風險是什麼？', 'Exact announcement timestamp、sector concentration、winner dependence 與真實成交可行性。'],
            ['為什麼 FRR 沒有升級？', '因為 remove top 10 winners 後轉負，而且高流動性版本變弱；它比較適合作為 S1 sizing diagnostic。'],
            ['下一步怎麼做？', '先補公司級 announcement timestamp 與 paper trading log，驗證資料更新與 next-open fill 是否符合回測假設。'],
        ]
        docx_table(doc, ['問題', '建議回答'], qas)


def build_docx(charts: list[tuple[str, str]], tables: dict[str, list[list[str]]]) -> None:
    d = Document(); add_docx_common(d, '台股月營收驚喜策略研究作品集', '履歷附件版 V3｜圖表強化與 Word/PDF 版本', False, charts, tables); d.save(RESUME_DOCX)
    g = Document(); add_docx_common(g, '台股月營收驚喜策略研究作品集', '面試講稿版 V3｜說法、追問與圖表講解', True, charts, tables); g.save(GUIDE_DOCX)


def make_styles():
    s = getSampleStyleSheet()
    s.add(ParagraphStyle('TitleC', fontName=FONT, fontSize=24, leading=31, textColor=NAVY, alignment=TA_CENTER, spaceAfter=10))
    s.add(ParagraphStyle('SubC', fontName=FONT, fontSize=11, leading=16, textColor=MUTED, alignment=TA_CENTER, spaceAfter=8))
    s.add(ParagraphStyle('H1C', fontName=FONT, fontSize=15, leading=20, textColor=NAVY, spaceBefore=8, spaceAfter=6))
    s.add(ParagraphStyle('H2C', fontName=FONT, fontSize=11, leading=15, textColor=BLUE, spaceBefore=5, spaceAfter=4))
    s.add(ParagraphStyle('BodyC', fontName=FONT, fontSize=9, leading=13.2, textColor=colors.HexColor('#111827'), wordWrap='CJK', spaceAfter=4))
    s.add(ParagraphStyle('SmallC', fontName=FONT, fontSize=7.6, leading=10.2, textColor=MUTED, wordWrap='CJK'))
    s.add(ParagraphStyle('CapC', fontName=FONT, fontSize=7.5, leading=9.8, textColor=MUTED, alignment=TA_CENTER, spaceAfter=7))
    return s

S = make_styles()

def esc(x: str) -> str:
    return str(x).replace('&','&amp;').replace('<','&lt;').replace('>','&gt;')

def P(text: str, style='BodyC') -> Paragraph:
    return Paragraph(esc(text), S[style])

def pdf_table(headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> Table:
    data = [[Paragraph(f'<b>{esc(h)}</b>', S['SmallC']) for h in headers]]
    for r in rows:
        data.append([Paragraph(esc(str(c)), S['SmallC']) for c in r])
    if widths is None:
        widths = [16.9/len(headers)*cm] * len(headers)
    t = Table(data, colWidths=widths, repeatRows=1)
    t.setStyle(TableStyle([
        ('BACKGROUND',(0,0),(-1,0),NAVY),('TEXTCOLOR',(0,0),(-1,0),WHITE),('ROWBACKGROUNDS',(0,1),(-1,-1),[WHITE,LIGHT]),
        ('BOX',(0,0),(-1,-1),0.5,BORDER),('INNERGRID',(0,0),(-1,-1),0.35,BORDER),('VALIGN',(0,0),(-1,-1),'TOP'),
        ('LEFTPADDING',(0,0),(-1,-1),4),('RIGHTPADDING',(0,0),(-1,-1),4),('TOPPADDING',(0,0),(-1,-1),5),('BOTTOMPADDING',(0,0),(-1,-1),5),
    ]))
    return t

def pdf_chart(path: str, caption: str, max_h=8.6*cm) -> list[Any]:
    p = Path(path)
    if not p.exists():
        return [P(f'缺圖：{path}', 'SmallC')]
    img = Image(str(p))
    scale = min(16.8*cm/img.imageWidth, max_h/img.imageHeight, 1)
    img.drawWidth = img.imageWidth * scale; img.drawHeight = img.imageHeight * scale
    return [img, Paragraph(esc(caption), S['CapC'])]

def section(title: str) -> list[Any]:
    return [Paragraph(esc(title), S['H1C']), HRFlowable(width='100%', thickness=.6, color=BORDER, spaceAfter=6)]

def build_pdf_one(path: Path, title: str, subtitle: str, is_guide: bool, charts: list[tuple[str, str]], tables: dict[str, list[list[str]]]) -> None:
    story: list[Any] = []
    story += [Spacer(1, 2.0*cm), Paragraph(esc(title), S['TitleC']), Paragraph(esc(subtitle), S['SubC']), Spacer(1, .4*cm)]
    story += [pdf_table(['文件定位','狀態'], [['研究作品集 / 履歷附件' if not is_guide else '面試講稿 / 私人準備版', 'Research / Paper-Trading Only；非投資建議；非實盤交易系統']], [7.0*cm,9.9*cm]), PageBreak()]
    story += section('01｜Executive Snapshot')
    story += [P('本作品集以台灣月營收公告制度為起點，建立 SUR-style 基本面 surprise 因子、短週期 portfolio backtest、robustness gates、execution realism 與 paper-trading schema。最終保留 S1 作為 portfolio-grade v0.1 研究候選。')]
    story += [pdf_table(['版本','Return','Sharpe','MDD','定位'], tables['headline'], [3.8*cm,2.3*cm,2.0*cm,2.0*cm,6.8*cm]), PageBreak()]
    story += section('02｜策略規格與研究設計')
    story += [P('因果鏈：月營收公告制度 → 持續性營收 surprise → 投資人預期逐步調整 → 公布後 10–20D repricing。'), pdf_table(['模組','設定','目的'], tables['spec'], [3.2*cm,5.0*cm,8.7*cm]), Spacer(1,6), pdf_table(['資料 / Gate','做法','目的'], tables['robustness'], [3.4*cm,6.0*cm,7.5*cm]), PageBreak()]
    story += section('03｜核心績效圖表')
    for i, (p, cap) in enumerate(charts[:2], 1):
        story += pdf_chart(p, cap, 8.2*cm)
    story += [PageBreak()]
    story += section('04｜搜尋紀律與 Robustness')
    for p, cap in charts[2:5]:
        story += pdf_chart(p, cap, 6.3*cm)
    story += [PageBreak()]
    story += section('05｜Walk-forward、Sector、Extension')
    for p, cap in charts[5:9]:
        story += pdf_chart(p, cap, 5.6*cm)
    story += [PageBreak()]
    story += section('06｜FRR 新策略：融資清洗反彈')
    story += [P('FRR 測試「強營收 surprise + 融資去槓桿 + 放量換手 / 不接刀」是否形成短線反彈。第一輪結果顯示 headline 有訊號，但不足以取代 S1。')]
    story += [pdf_table(['Variant','Delay','Hold','Return','Sharpe','MDD','Trades'], tables['frr_top'], [4.6*cm,1.4*cm,1.4*cm,2.0*cm,1.8*cm,2.0*cm,1.5*cm])]
    for p, cap in charts[9:]:
        story += pdf_chart(p, cap, 6.2*cm)
    story += [PageBreak()]
    story += section('07｜FRR 壓力測試與 Sector')
    story += [pdf_table(['Variant','Remove top N','Return','Sharpe','MDD'], tables['frr_remove'][:8], [5.7*cm,2.2*cm,2.4*cm,2.2*cm,2.4*cm]), Spacer(1,8), pdf_table(['Variant','Subset','Return','Sharpe','MDD'], tables['frr_sector'], [5.7*cm,3.0*cm,2.4*cm,2.2*cm,2.4*cm]), PageBreak()]
    story += section('08｜限制、Roadmap、最終定位')
    story += [pdf_table(['Gate','下一步','目的'], tables['roadmap'], [2.5*cm,6.4*cm,8.0*cm]), P('結論：S1 是目前主要 portfolio-grade research candidate；FRR 暫時只保留為 timing / sizing diagnostic。最專業的敘述不是「策略已可實盤」，而是展示完整研究流程與誠實的 promotion discipline。')]
    if is_guide:
        story += [PageBreak()] + section('09｜Interview Talking Guide')
        story += [Paragraph('30 秒版本', S['H2C']), P('我做的是台股月營收 surprise 的短週期量化研究。核心不是技術指標，而是利用台灣每月公布營收的制度，測試持續性營收驚喜是否會在 10–20 個交易日內被市場逐步重估。研究中我做了 OOS、sector survival、remove-winner、execution timing 與 paper-trading schema，所以最後我把 S1 定位為 portfolio-grade research candidate，而不是 production-ready alpha。')]
        story += [pdf_table(['問題','建議回答'], [['為什麼不用單月 YoY？','單月 YoY 容易受季節與一次性因素干擾，所以用 3M SUR persistence。'], ['最大風險？','Exact announcement timestamp、sector concentration、winner dependence、真實成交可行性。'], ['FRR 為什麼不升級？','Remove top 10 後轉負，高流動性版本變弱，因此只當 S1 diagnostic。'], ['下一步？','補 exact timestamp，做 3–6 個月 paper trading，驗證 non-fill / 滑價 / timing。']], [5.0*cm,11.9*cm])]
    doc = SimpleDocTemplate(str(path), pagesize=A4, leftMargin=1.7*cm, rightMargin=1.7*cm, topMargin=1.3*cm, bottomMargin=1.25*cm)
    doc.build(story, onFirstPage=footer, onLaterPages=footer)

def footer(canvas, doc):
    canvas.saveState(); canvas.setFont(FONT, 7); canvas.setFillColor(MUTED)
    canvas.drawString(1.7*cm, .75*cm, 'Taiwan Monthly Revenue Strategy Research｜Research / Paper-Trading Only')
    canvas.drawRightString(A4[0]-1.7*cm, .75*cm, f'Page {doc.page}')
    canvas.restoreState()

def build_pdfs(charts: list[tuple[str, str]], tables: dict[str, list[list[str]]]) -> None:
    build_pdf_one(RESUME_PDF, '台股月營收驚喜策略研究作品集', '履歷附件版 V3｜圖表強化與 Word/PDF 版本', False, charts, tables)
    build_pdf_one(GUIDE_PDF, '台股月營收驚喜策略研究作品集', '面試講稿版 V3｜說法、追問與圖表講解', True, charts, tables)

def main() -> int:
    extra = make_extra_charts()
    charts = existing_charts() + extra
    tables = collect_tables()
    build_docx(charts, tables)
    build_pdfs(charts, tables)
    for p in [RESUME_DOCX, RESUME_PDF, GUIDE_DOCX, GUIDE_PDF]:
        print(p, p.stat().st_size)
    return 0

if __name__ == '__main__':
    raise SystemExit(main())
