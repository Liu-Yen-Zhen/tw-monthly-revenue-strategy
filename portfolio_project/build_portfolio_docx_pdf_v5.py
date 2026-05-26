#!/usr/bin/env python3
"""Build V5 portfolio deliverables modeled after the uploaded ETF research PDF.

V5 keeps the monthly-revenue research content but mirrors the ETF portfolio style:
cover KPI strip, numbered sections, dense professional narrative, detailed tables,
chart captions, honest strategy feasibility conclusion, and future research roadmap.
"""
from __future__ import annotations

import csv, re, zipfile
from pathlib import Path
from typing import Any
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import Image, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle, HRFlowable

ROOT=Path('/Users/liuyenzhen/quant-research/tw_monthly_revenue')
PORT=ROOT/'portfolio_project'; CHARTS=ROOT/'reports'/'charts'; V3CHARTS=PORT/'charts_v3'; PROCESSED=ROOT/'data'/'processed'
RESUME_DOCX=PORT/'Taiwan_Monthly_Revenue_Strategy_Resume_Version_ZH_PRO_V5.docx'
RESUME_PDF=PORT/'Taiwan_Monthly_Revenue_Strategy_Resume_Version_ZH_PRO_V5.pdf'
GUIDE_DOCX=PORT/'Taiwan_Monthly_Revenue_Strategy_Talking_Guide_ZH_PRO_V5.docx'
GUIDE_PDF=PORT/'Taiwan_Monthly_Revenue_Strategy_Talking_Guide_ZH_PRO_V5.pdf'
pdfmetrics.registerFont(UnicodeCIDFont('STSong-Light')); FONT='STSong-Light'
NAVY=colors.HexColor('#0F172A'); BLUE=colors.HexColor('#1D4ED8'); GREEN=colors.HexColor('#047857'); AMBER=colors.HexColor('#B45309'); RED=colors.HexColor('#B91C1C'); MUTED=colors.HexColor('#64748B'); LIGHT=colors.HexColor('#F8FAFC'); BORDER=colors.HexColor('#CBD5E1'); WHITE=colors.white

def read_csv(p:Path)->list[dict[str,Any]]:
    if not p.exists(): return []
    with p.open(encoding='utf-8') as f: return list(csv.DictReader(f))
def flt(x):
    try: return float(x)
    except Exception: return 0.0
def pct(x): return f'{flt(x):.1%}'
def num(x): return f'{flt(x):.2f}'

def chart_list():
    xs=[
    (CHARTS/'s1_nav_drawdown_zh.png','圖 1：S1 NAV / Drawdown','S1 累積績效與回撤路徑；用於判斷策略是否只是單點跳升，或具有可觀察的持續風險報酬輪廓。'),
    (CHARTS/'s1_monthly_returns_is_oos_zh.png','圖 2：月報酬 IS/OOS','檢查樣本內與樣本外月份貢獻；避免績效完全由單一 regime 或單一月份支撐。'),
    (CHARTS/'signal_search_sharpe_mdd_scatter_zh.png','圖 3：搜尋空間 Sharpe vs MDD','顯示研究不是只挑最高 Sharpe，而是把回撤、可解釋性與後續壓力測試一起納入。'),
    (CHARTS/'top_variants_sharpe_vs_remove5_zh.png','圖 4：Top variants vs remove-top-5','檢查候選策略是否由少數最大贏家支撐，作為 promotion gate。'),
    (CHARTS/'remove_winners_sharpe_decay_zh.png','圖 5：Remove-winners decay','量化 right-tail dependence；這是目前不宣稱 production-ready 的核心原因之一。'),
    (CHARTS/'phase3_12_walkforward_oos_nav_zh.png','圖 6：Walk-forward OOS NAV','train-selected rule 沒有穩定打敗固定 S1，支持保留簡潔 incumbent。'),
    (CHARTS/'phase3_12_sector_survival_zh.png','圖 7：Sector survival','顯示電子 / 半導體供應鏈依賴，策略定位不應誇大成 broad-market alpha。'),
    (CHARTS/'phase3_17_quiet_digestion_nav_zh.png','圖 8：Quiet digestion','price-volume/K-line 僅作為基本面 surprise 後的消化狀態診斷，不作為獨立技術指標策略。'),
    (CHARTS/'phase3_18_dynamic_sizing_nav_zh.png','圖 9：Dynamic sizing','quiet boost 作為 sizing overlay 有小幅改善，但 drawdown 未改善到足以升級。'),
    (V3CHARTS/'phase4_1_frr_top_sharpe_zh.png','圖 10：FRR Top Sharpe','新增融資清洗策略最高 Sharpe 仍低於 S1，適合保留為診斷層。'),
    (V3CHARTS/'phase4_1_frr_remove_winner_zh.png','圖 11：FRR remove-winner','FRR 移除 top winners 後快速轉弱，不適合作為獨立主策略。'),
    (V3CHARTS/'phase4_1_frr_sector_zh.png','圖 12：FRR sector survival','FRR 有部分非半導體訊號，但樣本數太少，仍需更多 OOS。')]
    return [(str(p),t,c) for p,t,c in xs if p.exists()]

def tables():
    frr=sorted(read_csv(PROCESSED/'frr_margin_deleveraging_variants.csv'), key=lambda r: flt(r.get('sharpe_cash_counted')), reverse=True)[:8]
    rem=read_csv(PROCESSED/'frr_margin_deleveraging_remove_winners.csv')
    return {
    'progress':[
      ['Phase 1–2','建立月營收資料與 proxy backtest','確認月營收 surprise 具短線研究價值','完成'],
      ['Phase 3.6–3.12','SUR / short-horizon / walk-forward / sector survival','S1 成為 incumbent；電子/半導體依賴明確','完成'],
      ['Phase 3.13–3.18','price-volume / K-line / quiet digestion / sizing','quiet digestion 有診斷價值，但不升級','完成'],
      ['Phase 3.19–3.26','execution realism / exact timing / OHLC-limit parser','保守執行後績效下降，需 paper trading','完成'],
      ['Phase 4.1','FRR 融資清洗反彈','有訊號但右尾依賴；作為 S1 diagnostic','完成']],
    'headline':[
      ['S1 incumbent','3M SUR + no overheated momentum + 20D','+167.5%','≈2.40','-7.9%','保留主候選'],
      ['S1 fixed-20 comparator','簡潔固定持有 benchmark','+161.9%','1.55','-21.2%','基準比較'],
      ['Quiet boost','quiet digestion sizing overlay','+174.4%','1.62','-21.2%','診斷/加權候選'],
      ['Conservative execution','next-open + 1.0% cost + non-fill stress','+111.9%','1.24','-24.9%','保守可行性檢查'],
      ['FRR best first-pass','margin deleveraging + volume absorption','+251.4%','1.55','-17.6%','不升級，僅診斷']],
    'windows':[
      ['Signal window','月營收公布後可觀察之 signal date','避免使用尚未公開資料'],
      ['Entry','next-open / delay 1–3 trading days','保守處理公告時間與盤後資料'],
      ['Holding','10D / 15D / 20D；S1 主版本 20D','符合短線 repricing 假說'],
      ['Cost','0.7%–1.2%；核心保守版 1.0%','含稅費與滑價 proxy'],
      ['Non-fill','possible limit-up exclusion','台股漲停排隊風險不可忽略']],
    'risk':[
      ['Exact timing','缺公司級歷史公告 timestamp','目前不宣稱 same-day tradability','高'],
      ['Survivorship / universe','歷史 universe 與下市資料仍待補強','可能高估可交易結果','中-高'],
      ['Winner dependence','remove-winner 後 Sharpe 衰退','需要 position sizing 與風險上限','高'],
      ['Sector concentration','電子/半導體供應鏈依賴','定位需限縮，不稱 broad alpha','高'],
      ['Execution realism','開盤跳空、漲停 non-fill、流動性容量','需 paper trading 驗證','高'],
      ['Overfitting','多輪搜尋可能追逐樣本內 winners','保留 S1 簡潔 incumbent','中']],
    'frr_top':[[r['variant'].replace('_',' '),r['delay_trading_days'],r['holding_days'],pct(r['total_return']),num(r['sharpe_cash_counted']),pct(r['mdd']),r['trades']] for r in frr],
    'frr_rem':[[r['variant'].replace('_',' '),r['remove_top_n'],pct(r['total_return']),num(r['sharpe_cash_counted']),pct(r['mdd'])] for r in rem if r.get('variant') in {'frr2_no_catch_falling_knife','frr3_volume_absorption'}],
    'paper':[
      ['Daily data cut','記錄資料實際可得時間：營收、價格、融資、OHLC/limit','驗證 anti-look-ahead'],
      ['Signal log','signal_generated_at、候選名單、排序、排除原因','驗證策略流程可重現'],
      ['Execution check','next-open 是否可成交、漲停/跌停、成交金額、預估滑價','驗證可執行性'],
      ['Outcome review','5D/10D/20D PnL、benchmark excess、sector exposure','比較 paper vs backtest'],
      ['Monthly committee','remove-winner、sector、capacity、cost drift review','決定是否進入下一 gate']],
    'qa':[
      ['一句話介紹','我研究台灣月營收制度是否造成短週期 post-disclosure drift，並用業界式 robustness gates 檢查策略是否值得 paper trading。'],
      ['為什麼有 alpha？','台灣月營收是高頻基本面資訊，持續性 surprise 可能更新市場對供應鏈需求的預期，而市場反應不一定即時完成。'],
      ['為什麼不是 production-ready？','exact timestamp、survivorship、execution fill、sector concentration、winner dependence 都還沒完全通過。'],
      ['FRR 怎麼講？','它是另一個有因果邏輯的短線策略，但 remove-winner 脆弱，因此先當 S1 timing/sizing diagnostic。'],
      ['你最大的亮點？','不是只做出漂亮績效，而是能建立研究假說、驗證、壓力測試，並誠實拒絕未通過 gate 的版本。']]
    }

# DOCX
def shade(cell, fill):
    tc=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); tc.append(shd)
def style_doc(doc):
    sec=doc.sections[0]; sec.top_margin=Inches(.55); sec.bottom_margin=Inches(.55); sec.left_margin=Inches(.65); sec.right_margin=Inches(.65)
    for n in ['Normal','Heading 1','Heading 2','Heading 3']:
        doc.styles[n].font.name='Microsoft JhengHei'; doc.styles[n]._element.rPr.rFonts.set(qn('w:eastAsia'),'Microsoft JhengHei')
    doc.styles['Normal'].font.size=Pt(10); doc.styles['Heading 1'].font.size=Pt(17); doc.styles['Heading 1'].font.color.rgb=RGBColor(15,23,42); doc.styles['Heading 2'].font.size=Pt(12.5); doc.styles['Heading 2'].font.color.rgb=RGBColor(29,78,216)
def dtable(doc, headers, rows):
    tb=doc.add_table(rows=1, cols=len(headers)); tb.style='Table Grid'
    for i,h in enumerate(headers):
        c=tb.rows[0].cells[i]; c.text=h; shade(c,'0F172A')
        for p in c.paragraphs:
            for r in p.runs: r.font.bold=True; r.font.color.rgb=RGBColor(255,255,255)
    for row in rows:
        cells=tb.add_row().cells
        for i,v in enumerate(row): cells[i].text=str(v)
    doc.add_paragraph('')
def dchart(doc,path,title,cap):
    if Path(path).exists(): doc.add_picture(path, width=Inches(6.65))
    p=doc.add_paragraph(f'{title}｜{cap}'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs: r.font.size=Pt(8); r.font.color.rgb=RGBColor(100,116,139)
def doc_content(doc,title,sub,is_guide,ch,tb):
    style_doc(doc); h=doc.add_heading(title,0); h.alignment=WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(sub).alignment=WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('0050/高股息ETF 作品集版式參照｜Research / Paper-Trading Only｜非投資建議').alignment=WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('25 個月樣本｜S1 Sharpe≈2.40｜12 張圖表｜完整 promotion gates').alignment=WD_ALIGN_PARAGRAPH.CENTER; doc.add_page_break()
    doc.add_heading('一、研究概況',1)
    doc.add_paragraph('本研究以台灣上市櫃公司月營收公告制度為核心，檢驗「持續性營收驚喜」是否會在公布後 10–20 個交易日形成可觀察的短線 repricing。研究流程參照業界量化研究標準：先提出市場結構假說，再定義資料可得時間與反前視規則，接著進行因子設計、portfolio backtest、robustness gates、execution realism 與 paper-trading schema。')
    doc.add_paragraph('研究結論不是宣稱策略已可實盤，而是保留 S1 作為 portfolio-grade v0.1 research candidate。S1 具備清楚的基本面 surprise 邏輯與吸引人的 proxy 指標，但在 exact timestamp、survivorship、成交可行性與 paper trading 之前，仍不應稱為 production-ready alpha。')
    dtable(doc,['階段','內容','核心結果','狀態'],tb['progress']); dtable(doc,['策略/版本','設計','Return','Sharpe','MDD','狀態'],tb['headline'])
    doc.add_heading('二、主要發現：S1 具研究價值，但不是 broad-market production alpha',1)
    doc.add_paragraph('核心發現是：高 3M SUR 且股價尚未過熱的標的，在 20D 視窗有較佳的短線 repricing 表現。這支持「基本面資訊逐步被市場吸收」的假說。然而，sector survival 顯示績效偏向電子 / 半導體供應鏈；remove-winner stress 也顯示策略仍有 right-tail dependence。因此專業定位應是「台灣電子/半導體供應鏈月營收 surprise repricing candidate」，而非全市場通用 alpha。')
    for i,(p,t,c) in enumerate(ch[:7],1):
        dchart(doc,p,t,c)
        if i in {2,5}: doc.add_page_break()
    doc.add_heading('三、機制解釋：制度化月營收 → 預期修正 → 短線漂移',1)
    doc.add_paragraph('台灣月營收資料的特殊性在於，它比季報更高頻、比新聞更結構化，且能即時反映供應鏈訂單、價格與出貨變化。若市場參與者需要時間確認 surprise 的持續性，或資金進場受流動性與風險預算限制，則公布後的 repricing 可能不是一天完成，而是分布在 10–20D。S1 使用 3M SUR persistence，目的是降低單月雜訊並捕捉更穩定的 fundamental update。')
    dtable(doc,['視窗/假設','設計','目的'],tb['windows'])
    doc.add_heading('四、Extension Diagnostics：quiet digestion 與 FRR',1)
    doc.add_paragraph('後續 price-volume / K-line 研究不是技術指標 mining，而是把價格與成交量視為市場消化基本面 surprise 的狀態變數。quiet digestion 有直覺與局部改善，但未能通過 promotion gate。FRR 則測試「強基本面 + 融資去槓桿」的另一條短線邏輯：槓桿賣壓釋放後可能反彈；但第一輪仍呈現右尾依賴，因此不升級為獨立策略。')
    for p,t,c in ch[7:]: dchart(doc,p,t,c)
    dtable(doc,['FRR Variant','Delay','Hold','Return','Sharpe','MDD','Trades'],tb['frr_top']); dtable(doc,['FRR Variant','Remove top N','Return','Sharpe','MDD'],tb['frr_rem'][:8])
    doc.add_heading('五、策略可行性：成本、成交與誠實結論',1)
    doc.add_paragraph('保守 execution proxy 顯示，當進場從 close proxy 改為 next-open、成本提高至 1.0%，並排除可能漲停不可成交標的後，策略品質明顯下降。這不是失敗，而是業界研究中必要的 feasibility discount：一個策略如果只在理想價格、低成本、無 non-fill 下好看，不能直接推向 paper trading 或 live。')
    dtable(doc,['風險','目前狀態','影響','程度'],tb['risk'])
    doc.add_heading('六、Paper Trading Plan 與未來研究方向',1)
    doc.add_paragraph('下一步應停止單純追求更高 full-sample Sharpe，改以 paper-trading 驗證 operational feasibility。Paper trading 不用來證明 alpha，而是用來驗證資料是否準時可得、訊號是否可重現、實際開盤是否可成交、滑價與 non-fill 是否符合回測假設。')
    dtable(doc,['流程','記錄欄位/動作','驗證目的'],tb['paper'])
    doc.add_heading('七、研究紀律與最終定位',1)
    doc.add_paragraph('本作品集最有價值的地方，是展現研究紀律：從有因果邏輯的假說出發，保留簡潔 incumbent，對所有新增策略進行壓力測試，並誠實拒絕未通過 gate 的版本。對外建議說法：這是一個台股月營收 surprise 的 portfolio-grade quant research project，尚未是 production-ready live strategy。')
    if is_guide:
        doc.add_page_break(); doc.add_heading('八、面試講稿與答辯',1)
        doc.add_heading('30 秒版本',2); doc.add_paragraph('我做的是台股月營收 surprise 的短週期量化研究。核心從台灣每月公布營收的制度出發，測試持續性營收驚喜是否會在 10–20 個交易日內被市場逐步重估。我不只看 Sharpe，也做 OOS、sector survival、remove-winner、execution realism、融資策略延伸與 paper-trading schema，所以最後把 S1 定位為 portfolio-grade research candidate，而非 production-ready alpha。')
        doc.add_heading('3 分鐘版本',2); doc.add_paragraph('第一，市場結構：月營收提供高頻基本面資訊。第二，訊號設計：3M SUR persistence 避免單月雜訊，no-overheated momentum 避免已 pricing 標的。第三，投組：Top 8、industry cap、liquidity gate、20D horizon。第四，robustness：walk-forward、sector survival、remove-winner、cost、execution timing。第五，結論：S1 有研究價值但仍需 exact timestamp 與 paper trading；FRR 有邏輯但只當 diagnostic。')
        dtable(doc,['問題','回答'],tb['qa'])

# PDF
def make_styles():
    s=getSampleStyleSheet();
    s.add(ParagraphStyle('TitleC',fontName=FONT,fontSize=23,leading=30,textColor=NAVY,alignment=TA_CENTER,spaceAfter=8)); s.add(ParagraphStyle('SubC',fontName=FONT,fontSize=10.5,leading=15,textColor=MUTED,alignment=TA_CENTER,spaceAfter=6)); s.add(ParagraphStyle('H1C',fontName=FONT,fontSize=14.2,leading=19,textColor=NAVY,spaceBefore=7,spaceAfter=5)); s.add(ParagraphStyle('BodyC',fontName=FONT,fontSize=8.7,leading=12.7,textColor=colors.HexColor('#111827'),wordWrap='CJK',spaceAfter=4)); s.add(ParagraphStyle('SmallC',fontName=FONT,fontSize=7.1,leading=9.4,textColor=MUTED,wordWrap='CJK')); s.add(ParagraphStyle('CapC',fontName=FONT,fontSize=7.1,leading=9.3,textColor=MUTED,alignment=TA_CENTER,spaceAfter=5)); return s
S=make_styles()
def esc(x): return str(x).replace('&','&amp;').replace('<','&lt;').replace('>','&gt;')
def P(x,sty='BodyC'): return Paragraph(esc(x),S[sty])
def sect(x): return [Paragraph(esc(x),S['H1C']),HRFlowable(width='100%',thickness=.55,color=BORDER,spaceAfter=5)]
def ptable(headers,rows,widths=None):
    data=[[Paragraph(f'<b>{esc(h)}</b>',S['SmallC']) for h in headers]]+[[Paragraph(esc(c),S['SmallC']) for c in r] for r in rows]
    if widths is None: widths=[16.9/len(headers)*cm]*len(headers)
    t=Table(data,colWidths=widths,repeatRows=1); t.setStyle(TableStyle([('BACKGROUND',(0,0),(-1,0),NAVY),('TEXTCOLOR',(0,0),(-1,0),WHITE),('ROWBACKGROUNDS',(0,1),(-1,-1),[WHITE,LIGHT]),('BOX',(0,0),(-1,-1),.45,BORDER),('INNERGRID',(0,0),(-1,-1),.25,BORDER),('VALIGN',(0,0),(-1,-1),'TOP'),('LEFTPADDING',(0,0),(-1,-1),3.5),('RIGHTPADDING',(0,0),(-1,-1),3.5),('TOPPADDING',(0,0),(-1,-1),4),('BOTTOMPADDING',(0,0),(-1,-1),4)])); return t
def pchart(path,title,cap,maxh=6.6*cm):
    if not Path(path).exists(): return [P('缺圖 '+path,'SmallC')]
    img=Image(path); scale=min(16.8*cm/img.imageWidth,maxh/img.imageHeight,1); img.drawWidth=img.imageWidth*scale; img.drawHeight=img.imageHeight*scale
    return [img,Paragraph(esc(f'{title}｜{cap}'),S['CapC'])]
def footer(canvas,doc):
    canvas.saveState(); canvas.setFont(FONT,7); canvas.setFillColor(MUTED); canvas.drawString(1.55*cm,.72*cm,'Taiwan Monthly Revenue Surprise Research｜V5 Professional Portfolio｜Research Only'); canvas.drawRightString(A4[0]-1.55*cm,.72*cm,f'Page {doc.page}'); canvas.restoreState()
def build_pdf(path,title,sub,is_guide,ch,tb):
    story=[Spacer(1,1.5*cm),Paragraph(esc(title),S['TitleC']),Paragraph(esc(sub),S['SubC']),ptable(['S1 Sharpe','Total Return','MDD','圖表','狀態'],[['≈2.40','+167.5%','-7.9%','12','Research candidate']], [3.2*cm]*5),Spacer(1,.25*cm),P('Research / Paper-Trading Only｜非投資建議｜非實盤交易系統','SmallC'),PageBreak()]
    story+=sect('一、研究概況'); story+=[P('本研究以台灣上市櫃公司月營收公告制度為核心，檢驗「持續性營收驚喜」是否會在公布後 10–20 個交易日形成可觀察的短線 repricing。研究流程參照業界量化研究標準：市場結構假說、資料可得時間、反前視規則、因子設計、portfolio backtest、robustness gates、execution realism 與 paper-trading schema。'),P('結論不是宣稱策略已可實盤，而是保留 S1 作為 portfolio-grade v0.1 research candidate。'),ptable(['階段','內容','核心結果','狀態'],tb['progress'],[2.7*cm,5.0*cm,6.4*cm,2.0*cm]),PageBreak()]
    story+=sect('二、主要發現與 Headline Results'); story+=[ptable(['策略/版本','設計','Return','Sharpe','MDD','狀態'],tb['headline'],[3.2*cm,4.9*cm,1.8*cm,1.7*cm,1.7*cm,3.6*cm])]
    for p,t,c in ch[:2]: story+=pchart(p,t,c,6.9*cm)
    story+=[PageBreak()]
    story+=sect('三、Robustness：搜尋紀律、右尾依賴與 OOS');
    for p,t,c in ch[2:7]: story+=pchart(p,t,c,5.05*cm)
    story+=[PageBreak()]
    story+=sect('四、機制解釋與 Extension Diagnostics'); story+=[P('台灣月營收資料比季報更高頻、比新聞更結構化，能反映供應鏈訂單、價格與出貨變化。S1 使用 3M SUR persistence 來降低單月雜訊，並用 no-overheated momentum 避免已提前 pricing 的標的。')]
    for p,t,c in ch[7:9]: story+=pchart(p,t,c,5.7*cm)
    story+=[ptable(['視窗/假設','設計','目的'],tb['windows'],[3.2*cm,6.0*cm,7.7*cm]),PageBreak()]
    story+=sect('五、FRR 融資清洗反彈策略'); story+=[P('FRR 測試「強基本面 + 融資去槓桿」的短線邏輯：槓桿賣壓釋放後可能反彈。第一輪顯示放量換手版本較佳，但 remove-winner 脆弱、高流動性版本弱化，因此不升級為獨立策略。'),ptable(['FRR Variant','Delay','Hold','Return','Sharpe','MDD','Trades'],tb['frr_top'],[4.5*cm,1.2*cm,1.2*cm,1.9*cm,1.7*cm,1.9*cm,1.4*cm])]
    for p,t,c in ch[9:]: story+=pchart(p,t,c,5.2*cm)
    story+=[PageBreak()]
    story+=sect('六、策略可行性、風險與 Paper Trading Plan'); story+=[P('保守 execution proxy 顯示，當進場從 close proxy 改為 next-open、成本提高並排除可能漲停不可成交標的後，策略品質明顯下降。這是業界研究中必要的 feasibility discount。'),ptable(['風險','目前狀態','影響','程度'],tb['risk'],[3.0*cm,5.0*cm,6.3*cm,2.0*cm]),Spacer(1,6),ptable(['流程','記錄欄位/動作','驗證目的'],tb['paper'],[3.0*cm,7.0*cm,6.9*cm]),PageBreak()]
    story+=sect('七、研究紀律與最終定位'); story+=[P('本作品集最有價值的地方，是展現研究紀律：從有因果邏輯的假說出發，保留簡潔 incumbent，對所有新增策略進行壓力測試，並誠實拒絕未通過 gate 的版本。對外建議說法：這是一個台股月營收 surprise 的 portfolio-grade quant research project，尚未是 production-ready live strategy。')]
    if is_guide:
        story+=[PageBreak()]+sect('八、面試講稿與答辯'); story+=[P('30 秒版本：我做的是台股月營收 surprise 的短週期量化研究。核心從台灣每月公布營收的制度出發，測試持續性營收驚喜是否會在 10–20 個交易日內被市場逐步重估。我不只看 Sharpe，也做 OOS、sector survival、remove-winner、execution realism、融資策略延伸與 paper-trading schema，所以最後把 S1 定位為 portfolio-grade research candidate，而非 production-ready alpha。'),ptable(['問題','回答'],tb['qa'],[4.6*cm,12.3*cm])]
    doc=SimpleDocTemplate(str(path),pagesize=A4,leftMargin=1.55*cm,rightMargin=1.55*cm,topMargin=1.2*cm,bottomMargin=1.15*cm); doc.build(story,onFirstPage=footer,onLaterPages=footer)

def main():
    ch=chart_list(); tb=tables()
    d=Document(); doc_content(d,'台股月營收驚喜策略研究','履歷附件版 V5｜參照 ETF 作品集版式與敘事密度',False,ch,tb); d.save(RESUME_DOCX)
    g=Document(); doc_content(g,'台股月營收驚喜策略研究','面試講稿版 V5｜參照 ETF 作品集答辯風格',True,ch,tb); g.save(GUIDE_DOCX)
    build_pdf(RESUME_PDF,'台股月營收驚喜策略研究','履歷附件版 V5｜參照 ETF 作品集版式與敘事密度',False,ch,tb)
    build_pdf(GUIDE_PDF,'台股月營收驚喜策略研究','面試講稿版 V5｜參照 ETF 作品集答辯風格',True,ch,tb)
    for p in [RESUME_DOCX,RESUME_PDF,GUIDE_DOCX,GUIDE_PDF]: print(p,p.stat().st_size)
if __name__=='__main__': main()
